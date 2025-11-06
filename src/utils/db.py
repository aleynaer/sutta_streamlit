import os
import json
import docx
import streamlit as st
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore

# Initialize Firebase with your credentials
try:
    if not firebase_admin._apps:
        cred = credentials.Certificate({
            "type": st.secrets["firebase_credentials"]["type"],
            "project_id": st.secrets["firebase_credentials"]["project_id"],
            "private_key_id": st.secrets["firebase_credentials"]["private_key_id"],
            "private_key": st.secrets["firebase_credentials"]["private_key"],
            "client_email": st.secrets["firebase_credentials"]["client_email"],
            "client_id": st.secrets["firebase_credentials"]["client_id"],
            "auth_uri": st.secrets["firebase_credentials"]["auth_uri"],
            "token_uri": st.secrets["firebase_credentials"]["token_uri"],
            "auth_provider_x509_cert_url": st.secrets["firebase_credentials"]["auth_provider_x509_cert_url"],
            "client_x509_cert_url": st.secrets["firebase_credentials"]["client_x509_cert_url"],
            "universe_domain": st.secrets["firebase_credentials"]["universe_domain"]
        })
        firebase_admin.initialize_app(cred)

    db = firestore.client()
except Exception as e:
    st.error(f"""
        Failed to initialize Firebase: {str(e)}
        
        Please make sure:
        1. Firebase credentials are correct in .streamlit/secrets.toml
        2. Firestore API is enabled in Firebase Console
        3. Firebase project is properly set up
    """)
data_path = st.secrets["data_path"]


def init_db():
    collections = ['instructions', 'style_guides', 'dictionary', 'example_translations', 'user_translations']
    
    # Check if collections exist and initialize default data if needed
    for collection in collections:
        coll_ref = db.collection(collection)
        docs = coll_ref.limit(1).stream()
        if not list(docs):
            if collection == 'instructions':
                try:
                    # Read default instruction from DOCX
                    doc_path = os.path.join(data_path, "1- Talimat Dosyası.docx")
                    doc = docx.Document(doc_path)
                    talimat = '\n'.join([para.text for para in doc.paragraphs])
                    coll_ref.add({
                        'title': "Default Instructions",
                        'content': talimat,
                        'created_at': datetime.now()
                    })
                except Exception as e:
                    print(f"Error loading instructions: {str(e)}")
                    
            elif collection == 'style_guides':
                try:
                    # Read default style guide from DOCX
                    doc_path = os.path.join(data_path, "2-Cem Şen Çeviri Üslubu Rehberi.docx")
                    doc = docx.Document(doc_path)
                    uslup = '\n'.join([para.text for para in doc.paragraphs])
                    coll_ref.add({
                        'title': "Default Style Guide",
                        'content': uslup,
                        'created_at': datetime.now()
                    })
                except Exception as e:
                    print(f"Error loading style guide: {str(e)}")
                    
            elif collection == 'dictionary':
                try:
                    json_path = os.path.join(data_path, "sozluk.json")
                    with open(json_path, "r", encoding="utf-8") as f:
                        dictionary_data = json.load(f)
                        
                        if isinstance(dictionary_data, list):
                            for item in dictionary_data:
                                if isinstance(item, dict):
                                    pali = item.get('pali', '')
                                    turkish = item.get('turkish_translation', '')
                                    notes = item.get('explanation', '')
                                    if pali and turkish:
                                        coll_ref.add({
                                            'pali': pali,
                                            'turkish': turkish,
                                            'notes': notes,
                                            'created_at': datetime.now()
                                        })
                        else:
                            for pali_term, translations in dictionary_data.items():
                                if isinstance(translations, dict):
                                    turkish = translations.get('turkish_translation', '')
                                    notes = translations.get('explanation', '')
                                else:
                                    turkish = str(translations)
                                    notes = ''
                                    
                                coll_ref.add({
                                    'pali': pali_term,
                                    'turkish': turkish,
                                    'notes': notes,
                                    'created_at': datetime.now()
                                })
                                
                except Exception as e:
                    print(f"Error loading dictionary: {str(e)}")
                    raise
                    
            elif collection == 'example_translations':
                try:
                    txt_path = os.path.join(data_path, "ceviri_ornek.txt")
                    with open(txt_path, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                        
                        for line in lines:
                            parts = line.strip().split('|')
                            if len(parts) == 3:
                                title, original_text, translated_text = parts
                                coll_ref.add({
                                    'title': title,
                                    'original_text': original_text,
                                    'translated_text': translated_text,
                                    'is_selected': False,
                                    'created_at': datetime.now()
                                })
                except Exception as e:
                    print(f"Error loading example translations: {str(e)}")

# Get all instructions from the database
def get_all_instructions():
    instructions = []
    docs = db.collection('instructions').stream()
    for doc in docs:
        data = doc.to_dict()
        instructions.append((doc.id, data['title'], data['content']))
    return instructions

# Save and update instructions
def save_instruction(title, content):
    db.collection('instructions').add({
        'title': title,
        'content': content,
        'created_at': datetime.now()
    })

def update_instruction(id, title, content):
    db.collection('instructions').document(id).update({
        'title': title,
        'content': content
    })

# Get all style guides from the database
def get_all_style_guides():
    style_guides = []
    docs = db.collection('style_guides').stream()
    for doc in docs:
        data = doc.to_dict()
        style_guides.append((doc.id, data['title'], data['content']))
    return style_guides

# Save and update style guides
def save_style_guide(title, content):
    db.collection('style_guides').add({
        'title': title,
        'content': content,
        'created_at': datetime.now()
    })

def update_style_guide(id, title, content):
    db.collection('style_guides').document(id).update({
        'title': title,
        'content': content
    })

# Get all dictionary entries
def get_dictionary():
    """Get all dictionary entries for display"""
    dictionary = []
    docs = db.collection('dictionary').order_by('pali').stream()
    for doc in docs:
        data = doc.to_dict()
        dictionary.append((doc.id, data['pali'], data['turkish'], data.get('notes', '')))
    return dictionary

# Add, update, and delete dictionary entries
def add_dictionary_entry(pali, turkish, notes=""):
    db.collection('dictionary').add({
        'pali': pali,
        'turkish': turkish,
        'notes': notes,
        'created_at': datetime.now()
    })

def update_dictionary_entry(id, pali, turkish, notes):
    db.collection('dictionary').document(id).update({
        'pali': pali,
        'turkish': turkish,
        'notes': notes
    })

def delete_dictionary_entry(id):
    db.collection('dictionary').document(id).delete()

# Get dictionary as a dictionary object
def get_dictionary_as_dict():
    """Get dictionary as a Python dict for translation use"""
    dictionary = {}
    docs = db.collection('dictionary').stream()
    for doc in docs:
        data = doc.to_dict()
        dictionary[data['pali']] = {
            "turkish_translation": data['turkish'],
            "explanation": data.get('notes', '')
        }
    return dictionary

# Get all example translations
def get_all_examples():
    examples = []
    docs = db.collection('example_translations').stream()
    for doc in docs:
        data = doc.to_dict()
        examples.append((
            doc.id, 
            data['title'], 
            data['original_text'], 
            data['translated_text'], 
            data['is_selected'],
            data.get('created_at', datetime.now())
        ))
    return examples

# Save, update, delete, and select example translations
def save_example(title, original_text, translated_text):
    db.collection('example_translations').add({
        'title': title,
        'original_text': original_text,
        'translated_text': translated_text,
        'is_selected': False,
        'created_at': datetime.now()
    })

def update_example(id, title, original_text, translated_text):
    db.collection('example_translations').document(id).update({
        'title': title,
        'original_text': original_text,
        'translated_text': translated_text
    })

def delete_example(id):
    db.collection('example_translations').document(id).delete()

def update_example_selection(id, is_selected):
    db.collection('example_translations').document(id).update({
        'is_selected': is_selected
    })

def get_selected_examples():
    examples = []
    docs = db.collection('example_translations').where('is_selected', '==', True).stream()
    for doc in docs:
        data = doc.to_dict()
        examples.append((
            doc.id, 
            data['title'], 
            data['original_text'], 
            data['translated_text'],
            data['is_selected'],
            data.get('created_at', datetime.now())
        ))
    return examples

# Save user translations
def save_user_translation(title, original_text, translated_text):
    db.collection('user_translations').add({
        'title': title,
        'original_text': original_text,
        'translated_text': translated_text,
        'created_at': datetime.now()
    })

# Get user translations
def get_user_translations():
    translations = []
    docs = db.collection('user_translations').order_by('created_at', direction=firestore.Query.DESCENDING).stream()
    for doc in docs:
        data = doc.to_dict()
        translations.append((
            doc.id,
            data['title'],
            data['original_text'],
            data['translated_text'],
            data.get('created_at', datetime.now())
        ))
    return translations