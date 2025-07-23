from langchain_openai import ChatOpenAI
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
import os
import streamlit as st
import docx
import json
import pandas as pd
# Update imports
from utils.db import (init_db, get_all_instructions, save_instruction, update_instruction,
                     get_all_style_guides, save_style_guide, update_style_guide,
                     get_dictionary, add_dictionary_entry, update_dictionary_entry, 
                     delete_dictionary_entry, get_dictionary_as_dict,
                     get_all_examples, save_example, update_example, delete_example,
                     update_example_selection, get_selected_examples, get_user_translations,
                     save_user_translation)
# Initialize the database
init_db()

# Set up the OpenAI API key from Streamlit secrets
os.environ["OPENAI_API_KEY"] = st.secrets["openai_api_key"]

# Function to read text from a DOCX file
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

# Load instruction and style guide texts
#talimat = getText("src/data/1- Talimat Dosyası.docx")
#uslup = getText("src/data/2-Cem Şen Çeviri Üslubu Rehberi.docx")
def get_default_instruction():
    instructions = get_all_instructions()
    return instructions[0][2] if instructions else ""

def get_default_style_guide():
    style_guides = get_all_style_guides()
    return style_guides[0][2] if style_guides else ""

# Load example translation and dictionary content
data_folder = st.secrets["data_path"]
ceviri_ornek_path = os.path.join(data_folder, "ceviri_ornek.txt")
sozluk_path = os.path.join(data_folder, "sozluk.json")

# Load example translation text
with open(ceviri_ornek_path, "r") as ceviri_file:
    ceviri_content = ceviri_file.read()
# Load dictionary content
with open(sozluk_path, "r") as sozluk_file:
    sozluk_content = sozluk_file.read()


# Initialize the translation model
llm = ChatOpenAI(
    model="gpt-4.1-mini",
    temperature=0.7,
    max_tokens=10000,
    openai_api_key=os.getenv("OPENAI_API_KEY"),
)

# Function to translate text
def translate_text(llm, text, system_prompt):
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=text)
    ]
    response = llm.invoke(messages)
    return response.content, response.response_metadata.get("token_usage", {}).get("completion_tokens", 0), response.response_metadata.get("token_usage", {}).get("prompt_tokens", 0)

# Streamlit UI setup
st.title("Pali to Turkish Translation App")

# Sidebar for navigation
st.sidebar.title("Navigation")
options = st.sidebar.radio(
    "Select a page:", 
    ["Instructions", "Style Guide", "Dictionary", "Example Translations", "Translate", "My Translations"]
)

if options == "Instructions":
    st.header("Instructions")
    #st.write(talimat)
    # Add new instruction
    st.subheader("Add New Instruction")
    new_title = st.text_input("Title")
    new_content = st.text_area("Content")
    if st.button("Save New Instruction"):
        save_instruction(new_title, new_content)
        st.success("New instruction saved!")
        st.experimental_rerun()
    
    # Show existing instructions
    st.subheader("Existing Instructions")
    instructions = get_all_instructions()
    
    if instructions:
        selected_instruction = st.selectbox(
            "Select instruction",
            options=instructions,
            format_func=lambda x: x[1]  # Show title
        )
        
        if selected_instruction:
            edit_title = st.text_input("Edit Title", value=selected_instruction[1])
            edit_content = st.text_area("Edit Content", value=selected_instruction[2])
            
            if st.button("Update Instruction"):
                update_instruction(selected_instruction[0], edit_title, edit_content)
                st.success("Instruction updated!")
                st.experimental_rerun()
            
            # Use this instruction for translation
            if st.button("Use This Instruction"):
                st.session_state.current_instruction = edit_content
    else:
        st.info("No instructions saved yet. Please add a new instruction.")

elif options == "Style Guide":
    st.header("Style Guide")
    #st.write(uslup)
    # Add new style guide
    st.subheader("Add New Style Guide")
    new_title = st.text_input("Title", key="style_guide_title")
    new_content = st.text_area("Content", key="style_guide_content")
    if st.button("Save New Style Guide"):
        save_style_guide(new_title, new_content)
        st.success("New style guide saved!")
        st.experimental_rerun()
    
    # Show existing style guides
    st.subheader("Existing Style Guides")
    style_guides = get_all_style_guides()
    
    if style_guides:
        selected_style_guide = st.selectbox(
            "Select style guide",
            options=style_guides,
            format_func=lambda x: x[1],  # Show title
            key="style_guide_select"
        )
        
        if selected_style_guide:
            edit_title = st.text_input("Edit Title", value=selected_style_guide[1], key="edit_style_guide_title")
            edit_content = st.text_area("Edit Content", value=selected_style_guide[2], key="edit_style_guide_content")
            
            if st.button("Update Style Guide"):
                update_style_guide(selected_style_guide[0], edit_title, edit_content)
                st.success("Style guide updated!")
                st.experimental_rerun()
            
            # Use this style guide for translation
            if st.button("Use This Style Guide"):
                st.session_state.current_style_guide = edit_content
    else:
        st.info("No style guides saved yet. Please add a new style guide.")


elif options == "Dictionary":
    st.header("Dictionary")
    #st.json(sozluk_dict)
    # Add new entry section
    st.subheader("Add New Dictionary Entry")
    col1, col2, col3 = st.columns(3)
    with col1:
        new_pali = st.text_input("Pali Term")
    with col2:
        new_turkish = st.text_input("Turkish Translation")
    with col3:
        new_notes = st.text_input("Notes (Optional)")
    
    if st.button("Add Entry"):
        if new_pali and new_turkish:
            add_dictionary_entry(new_pali, new_turkish, new_notes)
            st.success("Entry added successfully!")
            st.experimental_rerun()
        else:
            st.error("Pali term and Turkish translation are required!")
    
    # Display and edit dictionary
    st.subheader("Dictionary Entries")
    dictionary = get_dictionary()
    
    if dictionary:
        # Create an editable dataframe
        edited_df = st.data_editor(
            pd.DataFrame(
                dictionary,
                columns=["ID", "Pali", "Turkish", "Notes"]
            ),
            key="dictionary_editor",
            num_rows="dynamic",
            hide_index=True,
            column_config={
                "ID": st.column_config.NumberColumn(
                    "ID",
                    help="Entry ID",
                    disabled=True,
                    width="small"
                ),
                "Pali": st.column_config.TextColumn(
                    "Pali Term",
                    width="medium"
                ),
                "Turkish": st.column_config.TextColumn(
                    "Turkish Translation",
                    width="medium"
                ),
                "Notes": st.column_config.TextColumn(
                    "Notes",
                    width="large"
                ),
                "Delete": st.column_config.CheckboxColumn(
                    "Delete",
                    help="Select to delete entry",
                    default=False,
                    width="small"
                )
            }
        )
        
        # Save changes button
        if st.button("Save Changes"):
            try:
                # Delete marked entries first
                for index, row in edited_df.iterrows():
                    if not pd.isna(row["ID"]) and row.get("Delete", False):
                        delete_dictionary_entry(int(row["ID"]))
                        continue
                    
                # Then process remaining entries
                for index, row in edited_df.iterrows():
                    if row.get("Delete", False):
                        continue
                    
                    if pd.isna(row["ID"]):  # New entry
                        if not pd.isna(row["Pali"]) and not pd.isna(row["Turkish"]):
                            add_dictionary_entry(
                                row["Pali"],
                                row["Turkish"],
                                row["Notes"] if not pd.isna(row["Notes"]) else ""
                            )
                    else:  # Update existing entry
                        update_dictionary_entry(
                            int(row["ID"]),
                            row["Pali"],
                            row["Turkish"],
                            row["Notes"] if not pd.isna(row["Notes"]) else ""
                        )
                st.success("Dictionary updated successfully!")
                st.experimental_rerun()
            except Exception as e:
                st.error(f"Error updating dictionary: {str(e)}")
    else:
        st.info("No dictionary entries found. Please add some entries.")


elif options == "Example Translations":
    st.header("Example Translations")
    st.write(example_text)
    # Add new example section
    st.subheader("Add New Example Translation")
    new_title = st.text_input("Title", key="example_title")
    col1, col2 = st.columns(2)
    with col1:
        new_original = st.text_area("Original Text", key="example_original")
    with col2:
        new_translation = st.text_area("Translation", key="example_translation")
    
    if st.button("Add Example"):
        if new_title and new_original and new_translation:
            save_example(new_title, new_original, new_translation)
            st.success("Example translation added successfully!")
            st.experimental_rerun()
        else:
            st.error("All fields are required!")
    
    # Display and edit examples
    st.subheader("Example Translations")
    examples = get_all_examples()
    selected_count = len([ex for ex in examples if ex[4]])  # Count selected examples
    
    if examples:
        for example in examples:
            with st.expander(f"Example: {example[1]}"):
                # Display original and translation
                cols = st.columns([3, 3, 1])
                with cols[0]:
                    st.text_area("Original", value=example[2], key=f"orig_{example[0]}", height=150)
                with cols[1]:
                    st.text_area("Translation", value=example[3], key=f"trans_{example[0]}", height=150)
                with cols[2]:
                    # Selection checkbox
                    is_selected = st.checkbox("Use as example", 
                                           value=example[4], 
                                           key=f"select_{example[0]}",
                                           disabled=selected_count >= 3 and not example[4])
                    
                    # Update selection
                    if is_selected != example[4]:
                        update_example_selection(example[0], is_selected)
                        st.experimental_rerun()
                    
                    # Edit and delete buttons
                    if st.button("Update", key=f"update_{example[0]}"):
                        update_example(
                            example[0],
                            example[1],
                            st.session_state[f"orig_{example[0]}"],
                            st.session_state[f"trans_{example[0]}"]
                        )
                        st.success("Example updated!")
                        st.experimental_rerun()
                    
                    if st.button("Delete", key=f"delete_{example[0]}"):
                        delete_example(example[0])
                        st.success("Example deleted!")
                        st.experimental_rerun()
    else:
        st.info("No example translations added yet.")

elif options == "My Translations":
    st.header("My Translations")
    translations = get_user_translations()
    
    if translations:
        for trans in translations:
            with st.expander(f"{trans[1]} - {trans[4]}"): # Title and date
                st.subheader("Original Text")
                st.write(trans[2])
                st.subheader("Translation")
                st.write(trans[3])
                
                # Download button for each translation
                doc = docx.Document()
                doc.add_heading(trans[1], 0)
                doc.add_heading("Original Text", level=1)
                doc.add_paragraph(trans[2])
                doc.add_heading("Turkish Translation", level=1)
                doc.add_paragraph(trans[3])
                
                # Save to temporary file
                doc_path = f"temp_{trans[1]}.docx"
                doc.save(doc_path)
                
                # Read file and provide download button
                with open(doc_path, "rb") as file:
                    st.download_button(
                        label="Download as DOCX",
                        data=file,
                        file_name=f"{trans[1]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # Remove temporary file
                os.remove(doc_path)
    else:
        st.info("No saved translations yet.")

elif options == "Translate":
    st.header("Translate Pali Text")
    
    # Input fields
    doc_title = st.text_input("Enter a title for your translation:")
    input_text = st.text_area("Enter Pali text to translate:")
    
    col1, col2, col3 = st.columns([1,1,1])
    with col1:
        translate_button = st.button("Translate")
    
    # Store translation in session state
    if translate_button:
        current_instruction = getattr(st.session_state, 'current_instruction', get_default_instruction())
        current_style_guide = getattr(st.session_state, 'current_style_guide', get_default_style_guide())
        current_dictionary = get_dictionary_as_dict()
        selected_examples = get_selected_examples()
        
        examples_text = ""
        if selected_examples:
            examples_text = "\n\n".join([
                f"Örnek {i+1}:\nOrijinal:\n{ex[2]}\nÇeviri:\n{ex[3]}"
                for i, ex in enumerate(selected_examples)
            ])
            examples_prompt = f"Çeviri tarzını benimsemek için aşağıdaki örnek çevirileri incele:\n{examples_text}"
        else:
            examples_prompt = f"Çeviri tarzını benimsemek için aşağıdaki örnek çeviriyi incele:\n{example_text}"

        system_prompt = f"""
        {current_instruction}
        
        {current_style_guide}
        
        {json.dumps(current_dictionary, ensure_ascii=False)}
        
        {examples_prompt}
        
        Şimdi, kullanıcının verdiği Pali metni Türkçeye çevir. Çıktında sadece çeviri metnini ver, başka hiçbir şey ekleme. 
        Çeviri metni Türkçe dil kurallarına uygun olmalıdır. Çevirilerinde anlaşılır ve akıcı bir dil kullanmalısın.
        Amacın Türkçe'ye çeviri yaparken Pali metninin anlamını doğru bir şekilde anlam kaybına uğratmadan yansıtmaktır.
        
        Orjinal Metin:
        """
        
        translated_text, completion_tokens, prompt_tokens = translate_text(llm, input_text, system_prompt)
        st.session_state.current_translation = translated_text
        st.session_state.original_text = input_text
    
    # Display translation if available
    if hasattr(st.session_state, 'current_translation'):
        st.subheader("Translated Text")
        st.write(st.session_state.current_translation)
        
        with col2:
            if st.button("Save Translation"):
                if doc_title:
                    save_user_translation(
                        doc_title,
                        st.session_state.original_text,
                        st.session_state.current_translation
                    )
                    st.success("Translation saved successfully!")
                else:
                    st.error("Please enter a title for your translation")
        
        with col3:
            if st.button("Download as DOCX"):
                if doc_title:
                    doc = docx.Document()
                    doc.add_heading(doc_title, 0)
                    
                    doc.add_heading("Original Text", level=1)
                    doc.add_paragraph(st.session_state.original_text)
                    
                    doc.add_heading("Turkish Translation", level=1)
                    doc.add_paragraph(st.session_state.current_translation)
                    
                    # Save to temporary file
                    doc_path = f"temp_{doc_title}.docx"
                    doc.save(doc_path)
                    
                    # Read file and provide download button
                    with open(doc_path, "rb") as file:
                        btn = st.download_button(
                            label="Download DOCX",
                            data=file,
                            file_name=f"{doc_title}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Remove temporary file
                    os.remove(doc_path)
                else:
                    st.error("Please enter a title for your translation")
